from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT.parent / "MultiBug-Agent项目讲解与面试题.docx"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, size=10.5, bold=False, color=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    size = {1: 17, 2: 14, 3: 12}.get(level, 11)
    color = {1: "1F4E79", 2: "2E74B5", 3: "1F4E79"}.get(level, "1F4E79")
    set_run_font(run, size=size, bold=True, color=color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=10)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, size=10)


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    for line in code.strip().splitlines():
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("333333")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].text = text
        shade_cell(hdr[idx], "D9EAF7")
        for p in hdr[idx].paragraphs:
            for r in p.runs:
                set_run_font(r, size=9.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            for p in cells[idx].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=9)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for name in ["List Bullet", "List Number"]:
        styles[name].font.name = "Microsoft YaHei"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        styles[name].font.size = Pt(10)
    return doc


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("MultiBug-Agent 项目讲解与面试题库")
    set_run_font(run, size=24, bold=True, color="1F4E79")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于多模态证据的软件缺陷定位与自动化测试生成平台")
    set_run_font(run, size=13, color="555555")

    add_para(doc, "")
    add_para(
        doc,
        "用途：秋招项目复盘、简历讲解、面试追问准备、后续升级路线规划。",
        size=10.5,
        bold=True,
        color="333333",
    )
    add_para(
        doc,
        "阅读方式：先看第 1-4 章理解项目整体，再看第 8 章面试题库，最后按第 9 章继续升级项目。",
        size=10.5,
    )


def add_project_overview(doc):
    add_heading(doc, "1. 项目整体介绍", 1)
    add_para(
        doc,
        "MultiBug-Agent 是一个面向 Web 应用测试场景的缺陷分析平台。它把页面截图、用户描述、前端日志、后端日志、代码片段和操作录屏等证据统一收集起来，自动生成缺陷摘要、复现步骤、根因分析、可疑代码位置、修复建议，以及 pytest / Playwright 自动化测试用例。",
    )
    add_heading(doc, "1.1 一句话介绍", 2)
    add_para(
        doc,
        "我做了一个多模态 UI 缺陷定位与自动化测试生成平台，可以根据截图、日志、代码和用户描述自动分析 Web 应用 Bug，并生成复现步骤、根因分析、可疑代码位置和自动化测试用例。",
        bold=True,
    )
    add_heading(doc, "1.2 这个项目解决什么问题", 2)
    add_bullets(
        doc,
        [
            "真实测试工作中，Bug 证据分散在截图、console、接口日志、后端日志、代码和用户描述中，人工整理成本高。",
            "普通自动化测试只能执行脚本，不能解释为什么失败；本项目强调“定位 + 解释 + 生成回归测试”。",
            "大模型项目容易变成聊天壳，本项目先建立可运行的规则引擎和代码检索闭环，再把 LLM 作为增强层。",
            "面试展示时可以直接选择示例场景，点击分析，展示 Bug 报告和测试脚本，比只讲模型训练更直观。",
        ],
    )
    add_heading(doc, "1.3 最终产出", 2)
    add_table(
        doc,
        ["输入材料", "系统处理", "输出结果"],
        [
            ["截图", "读取尺寸和页面状态线索，后续可接 VLM 生成视觉 caption", "页面异常线索、截图证据"],
            ["前后端日志", "正则提取错误类型、接口、堆栈、关键词", "缺陷类型、关键证据"],
            ["代码目录", "扫描代码文件，按函数/代码块切分并打分", "可疑文件、函数、代码片段"],
            ["用户描述", "补充业务上下文和复现动作", "复现步骤、影响模块"],
            ["录屏", "抽取关键帧，表示操作路径", "关键操作节点、页面变化"],
            ["分析结果", "规则推理 + 可选 DeepSeek 二次审查", "Bug 报告、pytest、Playwright"],
        ],
        widths=[1.25, 2.55, 2.55],
    )


def add_architecture(doc):
    add_heading(doc, "2. 系统架构和执行流程", 1)
    add_heading(doc, "2.1 总体流程", 2)
    add_numbers(
        doc,
        [
            "用户在 Streamlit 页面选择示例场景，或上传截图、日志、代码 zip、录屏和描述。",
            "app.py 把输入封装成 UploadedCase 数据对象。",
            "bug_reasoner.py 作为编排层，依次调用日志解析、截图分析、代码检索、缺陷推理和测试生成。",
            "log_parser.py 从日志中提取错误级别、错误类别、接口路径、堆栈文件和关键词。",
            "code_parser.py 扫描代码目录，把代码按函数或代码块切分为 CodeSnippet。",
            "code_locator.py 根据日志关键词和错误类别给代码片段打分，返回 Top-N 可疑位置。",
            "test_generator.py 按缺陷类型生成 pytest / Playwright 测试模板。",
            "report_generator.py 把所有结果渲染成 Markdown 报告，并保存到 outputs 目录。",
            "如果开启 AI 二次审查，multimodal_context.py 会把多模态证据组装成证据包，llm_client.py 调 DeepSeek API 做增强分析。",
        ],
    )
    add_heading(doc, "2.2 代码目录职责", 2)
    add_table(
        doc,
        ["文件/目录", "作用", "面试时怎么讲"],
        [
            ["app.py", "Streamlit 页面入口，负责上传、选择案例、触发分析、展示报告", "这是用户交互层，不放复杂业务逻辑。"],
            ["run_demo.py", "命令行入口，用于不启动页面时验证核心流程", "说明项目不是只能靠 UI 展示，核心能力可脚本化。"],
            ["src/models.py", "定义 UploadedCase、BugAnalysis、CodeHit 等数据结构", "用 dataclass 保证模块间传递结构清楚。"],
            ["src/log_parser.py", "解析前后端日志，识别错误类型和关键词", "用正则和规则先建立稳定基线。"],
            ["src/code_parser.py", "扫描代码文件，按函数/块切分", "为代码定位和后续向量检索打基础。"],
            ["src/code_locator.py", "根据关键词和错误类别给代码片段打分", "实现轻量代码检索，不依赖外部服务。"],
            ["src/bug_reasoner.py", "缺陷分析编排层", "核心流程入口，体现系统设计能力。"],
            ["src/test_generator.py", "按缺陷类型生成测试用例", "把 Bug 分析转成测试开发岗位关心的回归测试。"],
            ["src/report_generator.py", "生成 Markdown 报告和测试文件", "实现从分析到交付物的闭环。"],
            ["src/video_analyzer.py", "录屏关键帧抽取", "让项目具备真正的多模态扩展点。"],
            ["src/llm_client.py", "DeepSeek/OpenAI-compatible API 调用封装", "把模型接入和业务逻辑解耦。"],
            ["src/multimodal_context.py", "多模态证据包构造", "让模型看到结构化证据，而不是随便拼 prompt。"],
        ],
        widths=[1.6, 2.45, 2.45],
    )


def add_why_ui(doc):
    add_heading(doc, "3. 为什么能出现这个界面", 1)
    add_para(
        doc,
        "界面来自 Streamlit。Streamlit 是一个 Python Web 应用框架，它允许开发者用 Python 函数直接声明页面元素，例如标题、按钮、文件上传、图片展示、分栏和下载按钮。执行 streamlit run app.py 后，Streamlit 会启动本地 Web 服务，并在浏览器中渲染 app.py 中的组件。",
    )
    add_heading(doc, "3.1 app.py 中的关键 UI 代码含义", 2)
    add_table(
        doc,
        ["代码/组件", "作用", "显示效果"],
        [
            ["st.set_page_config(...)", "设置页面标题和布局", "浏览器页面使用宽屏布局"],
            ["st.title(...)", "显示主标题", "页面顶部出现 MultiBug-Agent"],
            ["st.sidebar", "定义左侧侧边栏区域", "输入方式、AI 开关、输出内容选项"],
            ["st.radio(...)", "单选输入方式", "示例场景 / 手动上传"],
            ["st.file_uploader(...)", "上传文件组件", "上传截图、录屏、代码 zip"],
            ["st.columns([0.42, 0.58])", "左右分栏布局", "左边输入，右边分析结果"],
            ["st.image(...)", "展示图片或 GIF", "展示 Bug 截图和操作录屏"],
            ["st.button(...)", "触发分析按钮", "点击后调用 analyze_case"],
            ["st.markdown(report)", "渲染 Markdown 报告", "右侧显示结构化 Bug 报告"],
            ["st.download_button(...)", "下载文件", "下载报告和测试脚本"],
        ],
        widths=[1.9, 2.2, 2.2],
    )
    add_heading(doc, "3.2 Streamlit 的运行机制", 2)
    add_bullets(
        doc,
        [
            "Streamlit 不是传统前后端分离框架，它会把 Python 脚本从上到下执行，并把 st.xxx 调用转换成页面组件。",
            "用户点击按钮或改变选择时，Streamlit 会重新运行脚本，因此代码要写成“根据当前状态渲染页面”。",
            "复杂业务逻辑不要堆在 app.py，应该放到 src/ 下的模块里，app.py 只负责输入输出和页面展示。",
            "本项目用左右分栏展示输入和输出，是为了让面试演示更直观。",
        ],
    )


def add_tech_stack(doc):
    add_heading(doc, "4. 用到的核心技术", 1)
    add_table(
        doc,
        ["技术", "本项目怎么用", "面试回答重点"],
        [
            ["Python", "主语言，负责文件处理、分析流程、报告生成", "工程能力，不只是写脚本。"],
            ["Streamlit", "构建可视化交互页面", "快速做 AI 工具原型，适合 Demo。"],
            ["dataclass", "定义日志信号、代码片段、分析结果等结构", "减少字典乱传，提高可维护性。"],
            ["pathlib", "跨平台路径处理", "比字符串拼路径更安全清晰。"],
            ["正则表达式", "识别日志中的错误、接口、堆栈、关键词", "适合日志结构化的第一层解析。"],
            ["Pillow", "读取截图尺寸、生成模拟 UI 截图", "提供图像基础处理能力。"],
            ["imageio", "生成 GIF / 处理录屏帧", "为视频关键帧分析做基础。"],
            ["pytest", "后端接口测试模板", "验证接口状态码、返回结构和异常分支。"],
            ["Playwright", "前端 UI 自动化测试模板", "验证页面跳转、按钮、loading、错误提示。"],
            ["requests", "调用 DeepSeek / OpenAI-compatible API", "不绑定某一家 SDK，接口更轻量。"],
            ["python-dotenv", "读取 .env 配置", "API Key 不写死进代码。"],
            ["Markdown", "生成 Bug 报告", "方便导出、预览和上传 GitHub。"],
            ["GitHub", "项目展示、版本管理、面试可信度", "README、截图、示例数据和运行方式很重要。"],
        ],
        widths=[1.35, 2.65, 2.2],
    )
    add_heading(doc, "4.1 多模态在本项目里的含义", 2)
    add_para(
        doc,
        "多模态不是只调用一个能看图的大模型，而是系统能融合不同类型的证据。本项目当前融合了文本描述、截图、日志、代码、结构化元数据和录屏关键帧。当前版本先把视觉证据结构化为截图尺寸、页面状态线索和关键帧路径；后续可以接入 VLM 把截图和关键帧转成更准确的视觉 caption。",
    )
    add_bullets(
        doc,
        [
            "文本模态：用户描述、Bug 描述、需求上下文。",
            "图像模态：页面截图、错误弹窗截图、白屏截图。",
            "视频模态：操作录屏，通过关键帧表达用户路径。",
            "日志模态：console 日志、后端 error 日志、接口日志。",
            "代码模态：前端/后端文件、函数、堆栈位置。",
            "结构化模态：metadata、缺陷类型、测试结果、报告字段。",
        ],
    )


def add_modules(doc):
    add_heading(doc, "5. 核心模块怎么完成的", 1)
    sections = [
        (
            "5.1 日志解析 log_parser.py",
            [
                "用 ERROR_PATTERNS 定义错误类型，例如 server_error、frontend_type_error、timeout、permission、upload_limit、data_mismatch。",
                "parse_logs 会逐行读取前端日志和后端日志，判断 severity、source、line_no 和 category。",
                "extract_keywords 会从用户描述、接口路径、堆栈文件和错误信息中提取关键词，供代码检索使用。",
                "优势是稳定、可解释；局限是规则需要维护，复杂自然语言日志可能漏判。",
            ],
        ),
        (
            "5.2 截图分析 screenshot_analyzer.py",
            [
                "使用 Pillow 打开截图，读取宽高和灰度统计信息。",
                "如果截图接近空白，会提示可能是白屏或渲染失败。",
                "示例场景里 metadata 还提供了 screen_state 和 ui_hint，模拟真实 VLM 识别结果。",
                "后续升级时，可以用视觉模型替换 metadata，让模型自动识别弹窗、loading、表单错误、白屏等 UI 状态。",
            ],
        ),
        (
            "5.3 代码切分 code_parser.py",
            [
                "递归扫描代码目录，只处理 .py、.js、.jsx、.ts、.tsx、.java、.go 等代码文件。",
                "跳过 node_modules、venv、__pycache__ 等无关目录。",
                "根据 def、class、function、const、export 等模式把文件切成代码片段。",
                "这样做可以减少模型或检索模块处理的上下文长度，提高定位效率。",
            ],
        ),
        (
            "5.4 可疑代码定位 code_locator.py",
            [
                "把日志关键词、接口路径、文件名、错误类型和代码片段文本进行匹配。",
                "命中接口路径、堆栈文件、业务关键词会加分。",
                "如果错误类别是 TypeError，会额外关注 undefined、null、response.data 等代码。",
                "如果错误类别是 timeout，会额外关注 timeout、retry、loading 等代码。",
                "最终按分数排序，输出 Top-N 可疑代码位置。",
            ],
        ),
        (
            "5.5 缺陷推理 bug_reasoner.py",
            [
                "这是项目的编排层，不直接做所有细节，而是调用日志、截图、代码检索和测试生成模块。",
                "它根据 dominant_category 判断缺陷类型，再生成严重程度、影响模块、复现步骤、根因分析和修复建议。",
                "新增版本支持 use_llm=True，此时会把多模态证据包交给 LLM 做二次审查。",
                "这种设计体现了可扩展性：规则引擎是底座，LLM 是增强，而不是全部依赖模型。",
            ],
        ),
        (
            "5.6 测试生成 test_generator.py",
            [
                "根据 bug_type 选择测试模板。",
                "后端 pytest 主要验证接口状态码、返回结构、异常输入和边界场景。",
                "前端 Playwright 主要验证页面操作、loading 是否释放、错误提示是否出现。",
                "面试时要说明：当前生成的是高质量模板，不是对任意项目 100% 可直接运行，真实落地需要结合项目选择器、接口 fixture 和 mock 数据。",
            ],
        ),
        (
            "5.7 DeepSeek 接入 llm_client.py",
            [
                "通过 requests 调用 OpenAI-compatible 的 /chat/completions 接口。",
                "配置从 .env 读取，不把 API Key 写死到代码。",
                "默认 provider 是 deepseek，base_url 是 https://api.deepseek.com，模型是 deepseek-v4-flash。",
                "没有 API Key 时自动跳过，保证项目离线也能演示。",
            ],
        ),
        (
            "5.8 多模态上下文 multimodal_context.py",
            [
                "把用户描述、截图线索、录屏关键帧、日志信号和可疑代码片段整理成 evidence packet。",
                "LLM 看到的是结构化证据，而不是一大坨混乱文本。",
                "后续如果接 VLM，只需要先把截图/关键帧 caption 放进 evidence packet，后面的缺陷推理链路不用重写。",
            ],
        ),
    ]
    for title, bullets in sections:
        add_heading(doc, title, 2)
        add_bullets(doc, bullets)


def add_deepseek(doc):
    add_heading(doc, "6. DeepSeek API 和多模态升级讲法", 1)
    add_para(
        doc,
        "根据 DeepSeek 官方文档，DeepSeek API 使用 OpenAI/Anthropic 兼容格式；OpenAI-compatible base_url 是 https://api.deepseek.com，Chat Completion 接口是 /chat/completions。官方文档当前列出的模型包括 deepseek-v4-flash 和 deepseek-v4-pro，旧的 deepseek-chat / deepseek-reasoner 标注为将在 2026/07/24 废弃。项目中默认使用 deepseek-v4-flash，并保留 thinking 开关配置。",
    )
    add_heading(doc, "6.1 接入步骤", 2)
    add_numbers(
        doc,
        [
            "在 DeepSeek 平台申请 API Key。",
            "在项目根目录复制 .env.example 为 .env。",
            "填写 DEEPSEEK_API_KEY、LLM_API_BASE 和 LLM_MODEL。",
            "命令行运行 python run_demo.py --case case_01_register_500 --use-llm。",
            "页面中打开“AI 二次审查”开关。",
        ],
    )
    add_code(
        doc,
        """
LLM_PROVIDER=deepseek
DEEPSEEK_API_KEY=你的APIKey
LLM_API_BASE=https://api.deepseek.com
LLM_MODEL=deepseek-v4-flash
DEEPSEEK_THINKING=disabled
""",
    )
    add_heading(doc, "6.2 面试时要注意的边界", 2)
    add_bullets(
        doc,
        [
            "不要说“DeepSeek 直接识别了截图”。当前接入的是 DeepSeek 文本 Chat API，项目把视觉证据结构化后交给它审查。",
            "如果要真正让模型看图，需要接 Qwen-VL、GPT-4o、Gemini、InternVL 等支持图像输入的 VLM。",
            "本项目的高级点不是“调用 API”，而是先做了证据抽取、代码检索、报告生成和测试生成，再让模型做二次审查。",
            "LLM 只作为增强层，底层规则流程仍可离线运行，这样可控、可解释、稳定。",
        ],
    )
    add_heading(doc, "6.3 后续真正多模态升级方案", 2)
    add_numbers(
        doc,
        [
            "视觉 caption 层：用 VLM 读取截图，输出页面类型、异常组件、错误提示、loading 状态、白屏概率。",
            "视频关键帧层：用 OpenCV/imageio 抽帧，再用 VLM 给每一帧生成操作描述。",
            "DOM 辅助层：如果能拿到 HTML/DOM，把按钮 id、表单字段、错误节点一起加入证据包。",
            "代码检索层：用 FAISS/Chroma + embedding 做语义检索，不只依赖关键词。",
            "Agent 层：让模型规划下一步需要的证据，例如要求补充接口返回、网络 HAR、用户角色、浏览器版本。",
            "自动复现层：用 Playwright 根据复现步骤自动操作页面，收集截图、console 和 trace。",
            "评估层：构造固定 Bug 数据集，统计缺陷分类准确率、Top-3 定位命中率、测试生成可用率。",
        ],
    )


def add_interview_questions(doc):
    add_heading(doc, "7. 面试官可能问的问题和参考回答", 1)
    categories = {
        "项目介绍类": [
            ("请你介绍一下这个项目。", "这是一个面向 Web 应用测试场景的多模态缺陷分析平台。输入包括截图、日志、代码、用户描述和录屏，输出包括缺陷类型、复现步骤、根因分析、可疑代码位置和 pytest / Playwright 测试用例。它的核心不是单纯调用大模型，而是先做日志解析、代码检索和规则推理，再可选接入 DeepSeek 做二次审查。"),
            ("为什么选这个项目做秋招项目？", "它贴近真实测试开发工作流，能覆盖 Python、自动化测试、日志分析、代码检索、大模型应用和工程化页面，比单纯 OCR 或聊天机器人更容易展示业务价值。"),
            ("这个项目最核心的亮点是什么？", "一是多模态证据融合，二是从 Bug 分析到测试用例生成的闭环，三是 LLM 只是增强层，底层规则流程可解释、可离线运行。"),
            ("你觉得这个项目和普通自动化测试有什么区别？", "普通自动化测试主要执行预先写好的用例；这个项目会根据缺陷证据生成复现步骤、定位可疑代码并生成测试模板，更偏 AI 测试分析平台。"),
            ("你项目里的多模态体现在哪里？", "文本描述、截图、录屏关键帧、前后端日志、代码片段和结构化元数据都进入同一个分析流程。"),
        ],
        "架构设计类": [
            ("整体架构怎么分层？", "分为交互层 app.py、数据结构层 models.py、证据解析层 log/screenshot/video/code parser、缺陷推理层 bug_reasoner、模型增强层 llm_client/multimodal_context、交付层 report/test_generator。"),
            ("为什么不把所有逻辑写在 app.py？", "app.py 只负责页面交互。如果业务逻辑都写进去会难测、难复用。拆到 src 模块后，命令行、测试和页面可以复用同一套核心流程。"),
            ("run_demo.py 有什么用？", "它是命令行入口，可以不依赖页面直接跑核心分析流程，方便测试、CI 和面试时快速验证。"),
            ("UploadedCase 和 BugAnalysis 为什么用 dataclass？", "它们是模块间传递的数据契约，比随手传字典更清晰，也便于补字段、测试和类型提示。"),
            ("这个系统有哪些输入输出边界？", "输入是用户描述、截图、日志、代码目录或 zip、录屏；输出是 Markdown 报告和测试脚本。中间不直接修改用户代码。"),
            ("如果上传 zip，有安全风险吗？", "有 zip slip 风险，所以 app.py 中 safe_extract_zip 会检查解压路径必须在目标目录内，避免恶意 zip 写到项目外。"),
        ],
        "Streamlit/UI 类": [
            ("为什么能出现网页界面？", "因为用 Streamlit。streamlit run app.py 后启动本地服务，st.title、st.sidebar、st.file_uploader、st.columns 等组件会被渲染成浏览器页面。"),
            ("Streamlit 的执行机制是什么？", "它会在用户交互后从上到下重新运行脚本，根据当前状态重新渲染页面。"),
            ("为什么用左右分栏？", "左侧放输入证据，右侧放分析结果，适合面试演示，也符合工具型产品的扫描习惯。"),
            ("手动上传模式怎么处理代码 zip？", "先保存上传文件到临时目录，再用 safe_extract_zip 解压，之后把 code_root 传给核心分析流程。"),
            ("页面和命令行结果是否一致？", "一致，因为它们都调用 analyze_case，只是输入来源不同。"),
        ],
        "日志解析类": [
            ("日志解析怎么做的？", "用正则匹配常见错误关键词，如 500、TypeError、timeout、403、Payload Too Large、amount mismatch，并提取接口路径、文件名和关键词。"),
            ("为什么不用大模型直接读日志？", "规则解析稳定、便宜、可解释，适合作为第一层；大模型适合作为二次审查，而不是完全替代确定性逻辑。"),
            ("dominant_category 是怎么判断缺陷类型的？", "根据日志信号类别计分，error 权重更高，分数最高的类别作为主缺陷类型。"),
            ("日志关键词怎么用于代码定位？", "extract_keywords 会提取接口路径、堆栈文件、错误类名和业务词，code_locator 再用这些词匹配代码片段。"),
            ("规则会不会漏判？", "会，所以项目保留 unknown 类别，并通过 AI 二次审查和后续模型增强降低漏判。"),
        ],
        "代码定位类": [
            ("代码定位怎么实现？", "先扫描代码目录，按函数或代码块切分，然后把日志关键词和代码片段匹配打分，最后输出 Top-N 可疑位置。"),
            ("为什么要切分代码？", "减少上下文长度，让定位粒度更细，也方便后续接向量检索或 LLM 分析。"),
            ("打分规则有哪些？", "关键词命中加分；接口路径、文件名、堆栈信息权重更高；不同错误类别有额外特征加分，例如 TypeError 关注 null/undefined，timeout 关注 loading/timeout。"),
            ("这种代码定位有什么局限？", "它偏关键词检索，语义理解弱；后续可以接 embedding、AST、调用链和真实堆栈行号。"),
            ("为什么报告里输出代码片段？", "面试官和开发人员需要看到证据，只有文件名不够，可疑片段能提高解释性。"),
        ],
        "测试生成类": [
            ("pytest 和 Playwright 区别是什么？", "pytest 主要做 Python 后端/接口测试；Playwright 做浏览器自动化，验证页面交互和 UI 状态。"),
            ("测试用例怎么生成？", "根据缺陷类型选择模板，例如 500 生成接口异常测试，timeout 生成 loading 释放测试，权限异常生成 403 测试。"),
            ("生成的测试能直接运行吗？", "在真实项目中需要结合项目的 client fixture、选择器、路由和 mock 数据调整。当前项目生成的是可读、可迁移的回归测试模板。"),
            ("为什么要生成测试？", "Bug 修复后需要回归验证，测试用例是把缺陷分析转化为质量保障资产。"),
            ("Playwright 测试重点验证什么？", "验证页面能否打开、用户操作后是否出现预期提示、loading 是否释放、错误边界是否显示。"),
        ],
        "DeepSeek/LLM 类": [
            ("DeepSeek 是怎么接入的？", "通过 requests 调用 OpenAI-compatible 的 /chat/completions 接口，配置从 .env 读取，默认 base_url 是 https://api.deepseek.com，模型是 deepseek-v4-flash。"),
            ("没有 API Key 项目还能跑吗？", "能。LLM 是可选增强，没有 Key 时使用规则引擎和代码检索，保证 Demo 稳定。"),
            ("为什么不用 openai SDK？", "requests 更轻量，也方便兼容 DeepSeek、OpenAI-compatible 代理和其他模型服务。后续如果需要流式输出或工具调用，可以再换 SDK。"),
            ("DeepSeek 能直接看图吗？", "当前官方 Chat Completion 接口是文本接口，本项目把视觉证据结构化后交给 DeepSeek。真正看图要接 VLM，如 Qwen-VL、GPT-4o、Gemini 或 InternVL。"),
            ("LLM 增强具体增强什么？", "它基于证据包做二次审查，判断规则结论是否可信，补充根因分析、缺失证据和回归测试建议。"),
            ("怎么避免大模型胡说？", "prompt 明确要求基于证据，不编造日志和代码；同时保留规则证据和可疑代码片段，报告中展示来源。"),
            ("thinking 模式有什么用？", "DeepSeek 官方支持 thinking 开关和 reasoning_effort。复杂分析可以打开，但成本和延迟更高，Demo 默认 disabled。"),
        ],
        "多模态升级类": [
            ("现在的多模态是真多模态吗？", "是多源证据融合，但视觉理解当前是结构化线索和示例元数据。严格来说，真正 VLM 识图是下一步升级。面试时要诚实区分。"),
            ("如何加入真正截图理解？", "接入支持 image input 的 VLM，把 screenshot.png 发送给模型，生成 UI 状态 caption，再写入 evidence packet。"),
            ("如何处理录屏？", "先抽关键帧，再对关键帧做视觉 caption，最后按时间顺序总结用户操作路径。"),
            ("为什么不直接把视频传给模型？", "视频成本高、上下文大，关键帧更可控，也更容易和日志时间点对齐。"),
            ("如果要做高级 Agent，下一步是什么？", "让 Agent 根据已有证据决定要不要继续采集 HAR、DOM、接口返回、Playwright trace，然后再生成最终报告。"),
        ],
        "工程化和安全类": [
            ("API Key 怎么保护？", "放在 .env，不提交 GitHub；仓库只提交 .env.example。"),
            ("为什么要有 .gitignore？", "避免上传 .venv、缓存、输出文件、API Key 和大体积文件。"),
            ("上传 GitHub 应该注意什么？", "上传源码、README、示例数据和截图，不上传真实日志、密钥、公司代码和隐私图片。"),
            ("项目如何测试？", "用 pytest 验证核心流程，命令行 Demo 验证完整输入输出，页面手动验证交互。"),
            ("项目有哪些异常处理？", "截图读取失败会返回提示，zip 解压会检查路径，LLM 未配置或调用失败会降级到规则模式。"),
            ("如果日志很多怎么办？", "需要做日志采样、按时间窗口聚合、错误聚类和摘要，避免把全部日志塞给模型。"),
            ("如果代码仓库很大怎么办？", "需要增量索引、缓存、embedding 向量库和按错误上下文召回 Top-K 片段。"),
        ],
        "项目不足和优化类": [
            ("当前项目最大不足是什么？", "当前视觉理解还没有接真实 VLM，代码定位还是关键词打分，测试生成是模板级。"),
            ("你会怎么继续优化？", "接 VLM 做截图/关键帧 caption，接 FAISS 做语义代码检索，接 Playwright trace 自动复现，加入评估集统计定位准确率。"),
            ("如何证明项目有效？", "构造固定 Bug 数据集，统计缺陷分类准确率、Top-3 定位命中率、报告字段完整率和测试模板可用率。"),
            ("如果面试官说这是套壳 AI 项目，你怎么回应？", "我会说明模型只是增强层；项目核心包括日志结构化、代码检索、缺陷规则推理、报告生成、测试生成和可视化交互。即使没有 API Key，核心流程也能跑。"),
            ("这个项目适合哪个岗位？", "测试开发、AI 测试、大模型应用、Python 开发、自动化测试和软件质量工程。"),
        ],
    }
    for cat, qa_list in categories.items():
        add_heading(doc, cat, 2)
        for q, a in qa_list:
            add_para(doc, "Q：" + q, bold=True, color="1F4E79")
            add_para(doc, "A：" + a)


def add_upgrade_plan(doc):
    add_heading(doc, "8. 接下来怎么把项目做得更高级", 1)
    add_table(
        doc,
        ["优先级", "升级方向", "具体做法", "面试价值"],
        [
            ["P0", "修正项目讲解和 README", "熟悉流程，能手动演示两个案例", "先保证讲得清楚。"],
            ["P1", "接 DeepSeek 二次审查", "配置 .env，打开 AI 二次审查", "从规则系统升级到 LLM 增强系统。"],
            ["P1", "接视觉模型", "用 Qwen-VL/GPT-4o 分析截图和关键帧", "真正多模态能力。"],
            ["P2", "代码向量检索", "用 embedding + FAISS/Chroma 检索代码片段", "比关键词检索更高级。"],
            ["P2", "Playwright 自动复现", "根据复现步骤自动操作页面并收集 trace", "从报告生成升级到自动复现 Agent。"],
            ["P2", "评估集", "构造 30-50 个 Bug 场景，统计准确率", "有指标，显得工程成熟。"],
            ["P3", "FastAPI 后端", "把核心分析做成接口，前端独立", "更像企业级服务。"],
            ["P3", "Docker/CI", "容器化部署，GitHub Actions 跑测试", "工程化完整度高。"],
        ],
        widths=[0.7, 1.35, 2.5, 2.0],
    )
    add_heading(doc, "8.1 你现在最该做的三件事", 2)
    add_numbers(
        doc,
        [
            "能从头演示：启动页面、选择 case_01、点击分析、解释报告字段和可疑代码。",
            "能回答：为什么不是普通自动化测试，为什么不是简单大模型套壳。",
            "把 DeepSeek API 配好，演示“规则分析 + AI 二次审查”的区别。",
        ],
    )
    add_heading(doc, "8.2 官方资料来源", 2)
    add_bullets(
        doc,
        [
            "DeepSeek Quick Start：https://api-docs.deepseek.com/",
            "DeepSeek Chat Completion API：https://api-docs.deepseek.com/api/create-chat-completion",
            "Streamlit 文档：https://docs.streamlit.io/",
            "Playwright Python 文档：https://playwright.dev/python/",
            "pytest 文档：https://docs.pytest.org/",
        ],
    )


def add_run_commands(doc):
    add_heading(doc, "9. 你在 VS Code 里应该怎么跑", 1)
    add_para(doc, "如果 PowerShell 不允许激活虚拟环境，可以不激活，直接使用 .venv 里的 python.exe。")
    add_code(
        doc,
        r"""
.\.venv\Scripts\python.exe run_demo.py --case case_01_register_500
.\.venv\Scripts\python.exe -m streamlit run app.py
""",
    )
    add_para(doc, "配置 DeepSeek 后运行：")
    add_code(
        doc,
        r"""
copy .env.example .env
# 在 .env 中填写 DEEPSEEK_API_KEY
.\.venv\Scripts\python.exe run_demo.py --case case_01_register_500 --use-llm
.\.venv\Scripts\python.exe -m streamlit run app.py
""",
    )


def main():
    doc = setup_doc()
    add_cover(doc)
    doc.add_page_break()
    add_project_overview(doc)
    add_architecture(doc)
    add_why_ui(doc)
    add_tech_stack(doc)
    add_modules(doc)
    add_deepseek(doc)
    add_interview_questions(doc)
    add_upgrade_plan(doc)
    add_run_commands(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
